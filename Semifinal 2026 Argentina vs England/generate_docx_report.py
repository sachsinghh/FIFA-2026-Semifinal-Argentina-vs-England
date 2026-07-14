"""
Generates a detailed Word (.docx) methodology report from prediction_results.json,
walking through the Elo/Davidson submodel, the Poisson expected-goals submodel,
and the 60/40 blend step by step with the actual computed numbers.
"""

import json
import math
import os
from datetime import datetime

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
RESULTS_PATH = os.path.join(BASE_DIR, "prediction_results.json")
OUT_PATH = os.path.join(BASE_DIR, "Argentina vs England - Prediction Methodology Report.docx")

SUBTITLE_TEXT = "Sachi Singh - FIFA World Cup Semifinal Analysis - EDUCATIONAL PURPOSES ONLY"

NAVY = RGBColor(0x1a, 0x2b, 0x4a)
ACCENT_A = RGBColor(0x2a, 0x78, 0xd6)   # Argentina blue
ACCENT_DRAW = RGBColor(0xb0, 0x76, 0x00)  # Draw (darkened for print contrast)
ACCENT_B = RGBColor(0xc0, 0x2a, 0x2a)   # England red
MUTED = RGBColor(0x60, 0x60, 0x60)


def set_cell_shading(cell, hex_color):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_color)
    cell._tc.get_or_add_tcPr().append(shd)


def add_footer(document, text):
    section = document.sections[0]
    footer = section.footer
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(8)
    run.font.color.rgb = MUTED
    run.italic = True


def add_heading(document, text, level=1):
    h = document.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = NAVY
    return h


def add_body(document, text, *, bold=False, italic=False, size=11, color=None, space_after=8):
    p = document.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    p.paragraph_format.space_after = Pt(space_after)
    return p


def add_formula(document, text):
    p = document.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.75)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(10.5)
    run.font.color.rgb = NAVY
    return p


def add_bullets(document, items):
    for item in items:
        p = document.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        run.font.size = Pt(11)


def add_table(document, headers, rows, col_widths=None, header_fill="1a2b4a"):
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = ""
        run = hdr_cells[i].paragraphs[0].add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor(0xff, 0xff, 0xff)
        run.font.size = Pt(10)
        set_cell_shading(hdr_cells[i], header_fill)

    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            run = cells[i].paragraphs[0].add_run(str(val))
            run.font.size = Pt(10)

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)

    document.add_paragraph().paragraph_format.space_after = Pt(4)
    return table


def pct(x):
    return f"{x*100:.1f}%"


def main():
    with open(RESULTS_PATH) as f:
        r = json.load(f)

    team_a = r["predictions"]["team_a"]
    team_b = r["predictions"]["team_b"]
    preds = r["predictions"]
    poisson = r["submodels"]["poisson"]
    elo = r["submodels"]["elo_davidson"]
    stats_a = r["team_stats"]["team_a"]
    stats_b = r["team_stats"]["team_b"]
    h2h = r["h2h_record"]
    meth = r["methodology"]
    venue = r.get("venue")

    # ---- Recompute the intermediate worked numbers so the walkthrough is exact ----
    elo_a, elo_b, nu = elo["elo_team_a"], elo["elo_team_b"], elo["davidson_nu"]
    pi_a = 10 ** (elo_a / 400.0)
    pi_b = 10 ** (elo_b / 400.0)
    sqrt_term = math.sqrt(pi_a * pi_b)
    denom = pi_a + pi_b + nu * sqrt_term

    pw_a, pw_b = stats_a["poisson_window"], stats_b["poisson_window"]
    league_avg = pw_a["league_avg_goals"]
    attack_a = pw_a["goals_for_avg"] / league_avg
    defense_a = pw_a["goals_against_avg"] / league_avg
    attack_b = pw_b["goals_for_avg"] / league_avg
    defense_b = pw_b["goals_against_avg"] / league_avg
    lambda_a = poisson["lambda_team_a"]
    lambda_b = poisson["lambda_team_b"]

    poisson_vec = (poisson["team_a_win"], poisson["draw"], poisson["team_b_win"])
    elo_vec = (elo["team_a_win"], elo["draw"], elo["team_b_win"])
    w_p, w_e = meth["blend_weights"]["poisson"], meth["blend_weights"]["elo"]
    raw_blend = (
        w_p * poisson_vec[0] + w_e * elo_vec[0],
        w_p * poisson_vec[1] + w_e * elo_vec[1],
        w_p * poisson_vec[2] + w_e * elo_vec[2],
    )
    raw_sum = sum(raw_blend)

    # ================= DOCUMENT =================
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    for section in doc.sections:
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)

    add_footer(doc, SUBTITLE_TEXT)

    # ---- Title page ----
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(60)
    run = title_p.add_run(f"{team_a} vs {team_b}")
    run.font.size = Pt(30)
    run.bold = True
    run.font.color.rgb = NAVY

    sub1 = doc.add_paragraph()
    sub1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub1.add_run("FIFA World Cup 2026 Semifinal — Hybrid Poisson + Elo Prediction Model")
    run.font.size = Pt(15)
    run.font.color.rgb = MUTED

    sub2 = doc.add_paragraph()
    sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub2.paragraph_format.space_before = Pt(18)
    run = sub2.add_run(SUBTITLE_TEXT)
    run.font.size = Pt(12)
    run.bold = True
    run.italic = True
    run.font.color.rgb = NAVY

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_before = Pt(30)
    meta_lines = [
        f"Kickoff: {r['match_date_local']}",
        f"Venue: {venue['city']}, {venue['country']} ({'neutral venue' if venue['neutral'] else 'home advantage applies'})" if venue else "",
        f"Data as of: {r['data_as_of']}",
        f"Report generated: {r['generated_at'][:19].replace('T', ' ')} UTC",
    ]
    for line in meta_lines:
        if not line:
            continue
        rp = meta.add_run(line + "\n")
        rp.font.size = Pt(10.5)
        rp.font.color.rgb = MUTED

    doc.add_page_break()

    # ---- Executive Summary ----
    add_heading(doc, "1. Executive Summary", level=1)
    add_body(
        doc,
        f"This report documents a hybrid statistical model — 60% weight on a Poisson expected-goals model "
        f"and 40% weight on an Elo rating model — used to project the outcome of {team_a} vs {team_b} "
        f"({r['match']}). Both submodels are built entirely from real historical international match data "
        f"(including this tournament's actual results through the quarterfinal stage, as published in the "
        f"source dataset as of {r['data_as_of']}), not estimated or hand-typed inputs."
    )

    add_table(
        doc,
        ["Outcome", "Probability"],
        [
            [f"{team_a} Win", pct(preds["team_a_win"])],
            ["Draw", pct(preds["draw"])],
            [f"{team_b} Win", pct(preds["team_b_win"])],
        ],
    )

    outcomes = {f"{team_a} Win": preds["team_a_win"], "Draw": preds["draw"], f"{team_b} Win": preds["team_b_win"]}
    most_likely = max(outcomes, key=outcomes.get)
    add_body(
        doc,
        f"Most likely outcome: {most_likely} ({pct(outcomes[most_likely])}). "
        f"Most likely scoreline (from the Poisson grid): {poisson['most_likely_score']} "
        f"({pct(poisson['top_scorelines'][0]['probability'])}).",
        bold=True,
    )
    add_body(
        doc,
        "This is a probabilistic projection produced ahead of kickoff, not a report of an actual result. "
        "See Section 6 (Limitations) for the assumptions behind these numbers.",
        italic=True, color=MUTED, size=10,
    )

    # ---- Data Source ----
    add_heading(doc, "2. Data Source", level=1)
    add_body(
        doc,
        f"All inputs are computed from a historical international football results dataset spanning "
        f"{meth['elo_history_start']} to {r['data_as_of']} ({meth['elo_n_matches_used']:,} completed matches). "
        f"The dataset is a continuously updated public record of match results (source: "
        f"martj42/international_results on GitHub) and, as of the data snapshot used here, already includes "
        f"this World Cup's real group-stage and knockout results for both teams up to the quarterfinal round. "
        f"Unplayed fixtures (including the {team_a} vs {team_b} semifinal itself) are excluded from all "
        f"averages and rating calculations, since they have no score yet."
    )

    # ---- Methodology Overview ----
    add_heading(doc, "3. Methodology Overview", level=1)
    add_body(
        doc,
        "Two independent submodels are computed from the same historical dataset, then combined:"
    )
    add_bullets(doc, [
        "Poisson submodel (60% weight): estimates each team's expected goals from attacking and "
        "defensive scoring rates, then derives win/draw/loss probabilities from the distribution of "
        "possible scorelines.",
        "Elo submodel (40% weight): tracks each team's skill rating match-by-match over the full "
        "historical record, then converts the two teams' rating gap into win/draw/loss probabilities "
        "using a tie-adjusted rating model (plain Elo alone only estimates win-or-not, not draws).",
    ])
    add_body(
        doc,
        "The two submodels' probability vectors are combined as a weighted average (60% Poisson / 40% Elo) "
        "and renormalized to sum to 100%. The full worked calculation for both submodels follows."
    )

    # ---- Section 4: Elo submodel ----
    add_heading(doc, "4. Elo Submodel", level=1)

    add_heading(doc, "4.1 Building the Elo ratings", level=2)
    add_body(
        doc,
        f"Every team's Elo rating starts at 1500 and is updated after each of the {meth['elo_n_matches_used']:,} "
        f"historical matches in chronological order, using the standard Elo update rule adapted for football "
        f"(goal-difference weighting, home-advantage adjustment):"
    )
    add_formula(doc, "Expected(home) = 1 / (1 + 10^((EloAway - (EloHome + HomeAdv)) / 400))")
    add_formula(doc, "GoalMultiplier = min(1 + (|GoalDiff| - 1) / 8, 1.75)   [1.0 if the match was a draw]")
    add_formula(doc, "EloNew = EloOld + K x GoalMultiplier x (ActualResult - Expected)")
    add_body(
        doc,
        f"Parameters used: K-factor = {meth['elo_k_factor']:.0f}, home-advantage bonus = "
        f"+{meth['elo_home_advantage']:.0f} rating points (applied only to historical matches that were not "
        f"played at a neutral venue — the {team_a} vs {team_b} semifinal itself is neutral, so no "
        f"home-advantage adjustment is applied to either team for this match)."
    )
    add_table(
        doc,
        ["Team", "Current Elo Rating"],
        [[team_a, f"{elo_a:.1f}"], [team_b, f"{elo_b:.1f}"]],
    )

    add_heading(doc, "4.2 Converting Elo into Win / Draw / Loss probabilities", level=2)
    add_body(
        doc,
        "A plain Elo rating gap only estimates the probability that one side does not lose — it has no "
        "native concept of a draw. To split the outcome into three probabilities, this model uses the "
        "Davidson (1970) tie-adjusted extension of the Bradley-Terry model, which adds one extra parameter, "
        "ν (\"nu\"), that controls how much probability mass goes to a draw:"
    )
    add_formula(doc, "piA = 10^(EloA / 400)          piB = 10^(EloB / 400)")
    add_formula(doc, "P(A wins) = piA / (piA + piB + ν x sqrt(piA x piB))")
    add_formula(doc, "P(B wins) = piB / (piA + piB + ν x sqrt(piA x piB))")
    add_formula(doc, "P(draw)   = ν x sqrt(piA x piB) / (piA + piB + ν x sqrt(piA x piB))")
    add_body(
        doc,
        f"The draw parameter ν was not guessed — it was calibrated by maximum-likelihood estimation "
        f"against the outcomes of all {meth['elo_n_matches_used']:,} historical matches (a grid search over "
        f"candidate ν values, picking the one that best explains the actual historical mix of home wins, "
        f"draws, and away wins given each match's pre-match Elo gap). This produced ν = {nu:.2f}."
    )
    add_body(doc, "Plugging in the current ratings:")
    add_formula(doc, f"piA = 10^({elo_a:.1f} / 400) = {pi_a:.3f}")
    add_formula(doc, f"piB = 10^({elo_b:.1f} / 400) = {pi_b:.3f}")
    add_formula(doc, f"sqrt(piA x piB) = {sqrt_term:.3f}")
    add_formula(doc, f"Denominator = {pi_a:.3f} + {pi_b:.3f} + {nu:.2f} x {sqrt_term:.3f} = {denom:.3f}")
    add_table(
        doc,
        ["Outcome", "Elo/Davidson formula result"],
        [
            [f"{team_a} win", f"{pi_a:.3f} / {denom:.3f} = {pct(elo['team_a_win'])}"],
            ["Draw", f"{nu:.2f} x {sqrt_term:.3f} / {denom:.3f} = {pct(elo['draw'])}"],
            [f"{team_b} win", f"{pi_b:.3f} / {denom:.3f} = {pct(elo['team_b_win'])}"],
        ],
    )

    # ---- Section 5: Poisson submodel ----
    add_heading(doc, "5. Poisson Submodel", level=1)

    add_heading(doc, "5.1 Attack and defense strength", level=2)
    add_body(
        doc,
        f"Each team's scoring and conceding rates are measured over a rolling {meth['poisson_window_years']}-year "
        f"window of recent matches (ending at the dataset's most recent played match, {r['data_as_of']}), so the "
        f"estimate reflects current squad strength rather than being diluted by decades-old results."
    )
    add_table(
        doc,
        ["Metric", team_a, team_b],
        [
            ["Matches in window", pw_a["n_matches"], pw_b["n_matches"]],
            ["Goals scored / match", f"{pw_a['goals_for_avg']:.3f}", f"{pw_b['goals_for_avg']:.3f}"],
            ["Goals conceded / match", f"{pw_a['goals_against_avg']:.3f}", f"{pw_b['goals_against_avg']:.3f}"],
        ],
    )
    add_body(doc, f"League-average goals per team per match over the same window: {league_avg:.4f}.")
    add_formula(doc, "AttackStrength = GoalsScoredPerMatch / LeagueAverageGoals")
    add_formula(doc, "DefenseStrength = GoalsConcededPerMatch / LeagueAverageGoals")
    add_table(
        doc,
        ["Metric", team_a, team_b],
        [
            ["Attack strength", f"{attack_a:.3f}", f"{attack_b:.3f}"],
            ["Defense strength", f"{defense_a:.3f}", f"{defense_b:.3f}"],
        ],
    )
    add_body(
        doc,
        "An attack strength above 1.0 means the team scores more than the league average; a defense "
        "strength below 1.0 means the team concedes fewer goals than the league average (i.e. a stronger defense)."
    )

    add_heading(doc, "5.2 Expected goals (λ)", level=2)
    add_body(
        doc,
        "Each team's expected goals for this match (λ, lambda) combines its own attack strength with the "
        "opponent's defense strength, scaled by the league-average goal rate. No home-advantage term is "
        "applied since the venue is neutral:"
    )
    add_formula(doc, f"lambda({team_a}) = Attack({team_a}) x Defense({team_b}) x LeagueAvg")
    add_formula(doc, f"lambda({team_a}) = {attack_a:.3f} x {defense_b:.3f} x {league_avg:.3f} = {lambda_a:.3f}")
    add_formula(doc, f"lambda({team_b}) = Attack({team_b}) x Defense({team_a}) x LeagueAvg")
    add_formula(doc, f"lambda({team_b}) = {attack_b:.3f} x {defense_a:.3f} x {league_avg:.3f} = {lambda_b:.3f}")
    add_body(
        doc,
        f"Interpretation: {team_a} is expected to score {lambda_a:.2f} goals and {team_b} is expected to "
        f"score {lambda_b:.2f} goals in this match, on average, under the model's assumptions."
    )

    add_heading(doc, "5.3 From λ to a scoreline probability grid", level=2)
    add_body(
        doc,
        "Each team's goal count is modeled as an independent Poisson-distributed random variable with "
        "mean λ. The probability that a team scores exactly k goals is:"
    )
    add_formula(doc, "P(goals = k) = e^(-lambda) x lambda^k / k!")
    add_body(
        doc,
        f"Multiplying {team_a}'s and {team_b}'s goal-count probabilities together for every combination "
        f"from 0-0 up to 10-10 produces a full scoreline probability grid. Summing the cells where "
        f"{team_a} scores more gives P({team_a} win); summing where the scores are equal gives P(draw); "
        f"summing where {team_b} scores more gives P({team_b} win). The resulting probabilities are "
        f"renormalized to sum to exactly 100%:"
    )
    add_table(
        doc,
        ["Outcome", "Poisson submodel result"],
        [
            [f"{team_a} win", pct(poisson["team_a_win"])],
            ["Draw", pct(poisson["draw"])],
            [f"{team_b} win", pct(poisson["team_b_win"])],
        ],
    )
    add_body(doc, "Top 3 most likely scorelines from the grid:")
    add_table(
        doc,
        ["Scoreline", "Probability"],
        [[s["score"], pct(s["probability"])] for s in poisson["top_scorelines"]],
    )

    # ---- Section 6: Blend ----
    add_heading(doc, "6. Combining the Submodels (60% Poisson / 40% Elo)", level=1)
    add_body(
        doc,
        "The two submodels' win/draw/loss vectors are combined with fixed weights and renormalized so the "
        "three probabilities sum to exactly 100%:"
    )
    add_formula(doc, f"Blended = {w_p:.1f} x Poisson + {w_e:.1f} x Elo,  then Blended = Blended / sum(Blended)")
    add_table(
        doc,
        ["Outcome", "Poisson (60%)", "Elo (40%)", "Weighted sum", "Blended (normalized)"],
        [
            [f"{team_a} win",
             pct(poisson_vec[0]), pct(elo_vec[0]),
             f"{raw_blend[0]:.4f}", pct(preds["team_a_win"])],
            ["Draw",
             pct(poisson_vec[1]), pct(elo_vec[1]),
             f"{raw_blend[1]:.4f}", pct(preds["draw"])],
            [f"{team_b} win",
             pct(poisson_vec[2]), pct(elo_vec[2]),
             f"{raw_blend[2]:.4f}", pct(preds["team_b_win"])],
        ],
        col_widths=[1.1, 1.1, 1.0, 1.1, 1.2],
    )
    add_body(
        doc,
        f"(The weighted sums total {raw_sum:.4f} before renormalization, purely due to each submodel's own "
        f"internal rounding; dividing each by {raw_sum:.4f} yields the final blended percentages above, which "
        f"sum to exactly 100%.)",
        italic=True, size=10, color=MUTED,
    )

    add_heading(doc, "6.1 Final Result", level=2)
    final_table_rows = [
        [f"{team_a} win", pct(preds["team_a_win"])],
        ["Draw", pct(preds["draw"])],
        [f"{team_b} win", pct(preds["team_b_win"])],
    ]
    add_table(doc, ["Outcome", "Final Blended Probability"], final_table_rows)
    add_body(
        doc,
        f"Final prediction: {most_likely} is the most probable outcome at {pct(outcomes[most_likely])}, "
        f"with a most-likely scoreline of {poisson['most_likely_score']}.",
        bold=True,
    )

    # ---- Section 7: Supporting context ----
    add_heading(doc, "7. Supporting Context (not used in the probability calculation)", level=1)
    add_body(
        doc,
        "The following figures are shown for narrative context only. They are not fed into either "
        "submodel or the blend above."
    )
    add_heading(doc, "7.1 Head-to-Head Record", level=2)
    a_key, b_key = f"{team_a.lower()}_wins", f"{team_b.lower()}_wins"
    add_table(
        doc,
        [f"{team_a} wins", "Draws", f"{team_b} wins", "Total meetings"],
        [[h2h[a_key], h2h["draws"], h2h[b_key], h2h["total_meetings"]]],
    )
    if h2h.get("last_match"):
        lm = h2h["last_match"]
        add_body(
            doc,
            f"Last meeting: {lm['date']} — {lm['home_team']} {lm['home_score']}-{lm['away_score']} "
            f"{lm['away_team']} ({lm['tournament']}).",
            size=10,
        )
    add_body(
        doc,
        "Note: the underlying dataset has no penalty-shootout field, so any past meeting decided on "
        "penalties (e.g. the 1998 World Cup last-16 tie) is recorded as a draw.",
        italic=True, size=10, color=MUTED,
    )

    add_heading(doc, "7.2 Recent Form (last 10 matches)", level=2)
    fa, fb = stats_a["form"], stats_b["form"]
    add_table(
        doc,
        ["Team", "Points per game", "Goals for / match", "Goals against / match"],
        [
            [team_a, f"{fa['ppg']:.2f}", f"{fa['gf_avg']:.2f}", f"{fa['ga_avg']:.2f}"],
            [team_b, f"{fb['ppg']:.2f}", f"{fb['gf_avg']:.2f}", f"{fb['ga_avg']:.2f}"],
        ],
    )

    # ---- Section 8: Limitations ----
    add_heading(doc, "8. Limitations and Caveats", level=1)
    add_bullets(doc, r["caveats"])

    # ---- Section 9: Parameter reference ----
    add_heading(doc, "9. Parameter Reference", level=1)
    add_table(
        doc,
        ["Parameter", "Value"],
        [
            ["Blend weights (Poisson / Elo)", f"{w_p*100:.0f}% / {w_e*100:.0f}%"],
            ["Poisson attack/defense window", f"{meth['poisson_window_years']} years"],
            ["Elo K-factor", f"{meth['elo_k_factor']:.0f}"],
            ["Elo home-advantage bonus", f"+{meth['elo_home_advantage']:.0f} points (non-neutral matches only)"],
            ["Elo history start date", meth["elo_history_start"]],
            ["Historical matches used", f"{meth['elo_n_matches_used']:,}"],
            ["Davidson draw parameter (nu)", f"{nu:.2f} (MLE-calibrated)"],
            ["Data as of", r["data_as_of"]],
        ],
    )

    doc.add_paragraph()
    closing = doc.add_paragraph()
    closing.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = closing.add_run(SUBTITLE_TEXT)
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = MUTED

    doc.save(OUT_PATH)
    print(f"[OK] Saved {OUT_PATH}")


if __name__ == "__main__":
    main()
